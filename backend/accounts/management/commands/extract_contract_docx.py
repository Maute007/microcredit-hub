"""Extrai texto de um ficheiro .docx para colar em contract_general_clauses ou gravar no singleton."""

from pathlib import Path

from django.core.management.base import BaseCommand


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            rowt = " | ".join(x for x in cells if x)
            if rowt:
                parts.append(rowt)
    return "\n".join(parts)


class Command(BaseCommand):
    help = (
        "Lê um .docx e imprime o texto extraído (parágrafos e tabelas). "
        "Use --save para gravar em SystemSettings.contract_general_clauses. "
        "Documentos só com páginas digitalizadas como imagem não têm texto — use OCR ou reescreva no Word."
    )

    def add_arguments(self, parser):
        parser.add_argument(
            "path",
            nargs="?",
            default="docs/Contrato.docx",
            help="Caminho para o .docx (relativo à pasta backend/ por omissão).",
        )
        parser.add_argument(
            "--save",
            action="store_true",
            help="Grava o texto em SystemSettings.contract_general_clauses (singleton).",
        )

    def handle(self, *args, **options):
        raw = options["path"]
        p = Path(raw)
        if not p.is_file():
            base = Path(__file__).resolve().parents[3]
            p = base / raw
        if not p.is_file():
            self.stderr.write(self.style.ERROR(f"Ficheiro não encontrado: {raw}"))
            return

        text = _extract_docx_text(p)
        n = len(text)
        if n == 0:
            self.stdout.write(
                self.style.WARNING(
                    "Nenhum texto encontrado. O .docx pode ser só imagens (digitalização). "
                    "Converta para texto no Word ou use OCR, depois volte a exportar."
                )
            )
            return

        self.stdout.write(text)
        self.stdout.write("")
        self.stdout.write(self.style.SUCCESS(f"Total: {n} caracteres."))

        if options["save"]:
            from accounts.models import SystemSettings

            s = SystemSettings.get_solo()
            s.contract_general_clauses = text
            s.save(update_fields=["contract_general_clauses"])
            self.stdout.write(self.style.SUCCESS("Gravado em SystemSettings.contract_general_clauses."))
